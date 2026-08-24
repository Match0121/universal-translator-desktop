# -*- coding: utf-8 -*-
"""PDF 翻译后端（文本型 + OCR 图片型 + 提取方式可选）
能力：
  - extract_pdf：解析 PDF → 分类（text / image / empty）→ block 级提取段落（含标题层级）
  - start_ocr / ocr_status：逐页渲染 + RapidOCR 识别（后台线程 + 进度文件），文字层乱码时可强制使用
  - export_translated：原文段落 + 译文 → 生成 docx / md / txt / 重排 PDF（按标题层级分层）

依赖：pymupdf（import pymupdf）、python-docx（import docx）
      rapidocr_onnxruntime（仅 OCR 需要，缺失时该类 PDF 报错提示）
"""
import io
import json
import os
import threading
import unicodedata
from collections import Counter

import pymupdf
from docx import Document
from docx.shared import Pt

try:
    from rapidocr_onnxruntime import RapidOCR
    OCR_AVAILABLE = True
except Exception:
    RapidOCR = None
    OCR_AVAILABLE = False

# 文字层判定阈值：整 PDF 可提取字符数 > 50 视为文本型
TEXT_TYPE_THRESHOLD = 50
# OCR 渲染分辨率（dpi）：150 足够扫描件识别，速度约为 200 的 2 倍
OCR_DPI = 150
# 标题判定：字号 ≥ 正文基准 × 此系数，且文本较短
TITLE_SIZE_RATIO = 1.25
TITLE_MAX_LEN = 40


def _is_cjk(ch):
    return '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef'


def _is_noise(t):
    if not t or len(t) > 800:
        return True
    if t.isdigit() or all(ch in " .,%+-*/=<>:;|!?()[]{}'\"\`~^#$@&" for ch in t):
        return True
    return False


class PdfError(Exception):
    pass


def _main_font():
    """返回 (fontname, fontfile, measure_font)
    优先系统微软雅黑（中英一体、字形现代），不可用时退回内置 CJK 字体。
    """
    for p in (r'C:\Windows\Fonts\msyh.ttc', r'C:\Windows\Fonts\msyhbd.ttc'):
        if os.path.isfile(p):
            try:
                return 'ut-main', p, pymupdf.Font(fontfile=p)
            except Exception:
                continue
    return 'china-s', None, pymupdf.Font('china-s')


MAIN_FONT_NAME, MAIN_FONT_FILE, _MEASURE_FONT = _main_font()


def extract_pdf(data):
    """解析 PDF 字节，返回 {pdf_kind, total_chars, paragraphs, page_count}
    pdf_kind:
      'text'  - 有文字层，走正常提取
      'image' - 无文字层但有图像（扫描件/照片），需 OCR
      'empty' - 既无文字也无图像（空白文档）
    paragraphs: [{seq, original, title}]（title=True 表示该段为标题，字号显著大于正文）
    """
    try:
        doc = pymupdf.open(stream=data, filetype='pdf')
    except Exception as e:
        raise PdfError('无法解析 PDF: %s' % e)
    try:
        # 收集所有文本行（含坐标与字号），按页内 y 顺序聚类成段
        rows = []            # {page, y0, y1, size, text}
        has_images = False
        for page in doc:
            if page.get_images(full=True):
                has_images = True
            d = page.get_text('dict')
            for block in d.get('blocks', []):
                if block.get('type', 0) != 0:
                    continue            # 非文本块（图片块）跳过
                for line in block.get('lines', []):
                    spans = line.get('spans', [])
                    if not spans:
                        continue
                    parts = []
                    max_size = 0.0
                    for span in spans:
                        parts.append(span.get('text', ''))
                        max_size = max(max_size, span.get('size', 0) or 0)
                    text = ''.join(parts).strip()
                    if not text:
                        continue
                    x0, y0, x1, y1 = line.get('bbox', (0, 0, 0, 0))
                    rows.append({'page': page.number, 'x0': x0, 'y0': y0,
                                 'x1': x1, 'y1': y1, 'size': max_size, 'text': text})

        def _can_merge(last_row, row):
            # 跨页断开；行间空隙大于阈值视为段落分隔；字号突变（标题/强调）断开
            gap = row['y0'] - last_row['y1']
            line_h = max(last_row['size'] * 1.4, 10)
            if gap > max(line_h * 1.5, 8):
                return False
            if abs(row['size'] - last_row['size']) > 1.5:
                return False
            return True

        clusters = []        # [{lines: [...], page, rects, size}]
        for row in rows:
            if clusters and _can_merge(clusters[-1]['lines'][-1], row):
                clusters[-1]['lines'].append(row)
            else:
                clusters.append({'lines': [row]})
        for c in clusters:
            c['page'] = c['lines'][0]['page']
            c['rects'] = [[r['x0'], r['y0'], r['x1'], r['y1']] for r in c['lines']]

        sizes = []           # 全部字号（算正文基准用）
        total_chars = 0
        para_items = []      # [{text, size, page, rects}]
        for c in clusters:
            # 合并段内行：CJK 连续拼接，其他情况补空格（消断行）
            text = ''
            for row in c['lines']:
                lt = row['text']
                if not lt:
                    continue
                if text and not (_is_cjk(text[-1]) or _is_cjk(lt[0])):
                    text += ' '
                text += lt
            seg = max(r['size'] for r in c['lines'])
            sizes.append(seg)
            total_chars += len(text)
            para_items.append({'text': text, 'size': seg, 'page': c['page'],
                               'rects': c['rects']})
        # 正文基准字号：出现次数最多且字号较小者优先（并列时取小，避免标题字号抢基准）
        if sizes:
            cnt = Counter(sizes)
            base_size = min(cnt, key=lambda s: (-cnt[s], s))
        else:
            base_size = 12
        paras = []
        seq = 0
        for b in para_items:
            t = b['text'].strip()
            if _is_noise(t):
                continue
            is_title = b['size'] >= base_size * TITLE_SIZE_RATIO and len(t) <= TITLE_MAX_LEN
            paras.append({'seq': seq, 'original': t, 'title': is_title,
                          'page': b['page'], 'rects': b['rects']})
            seq += 1
        if total_chars > TEXT_TYPE_THRESHOLD:
            pdf_kind = 'text'
        elif has_images:
            pdf_kind = 'image'
        else:
            pdf_kind = 'empty'
        return {
            'pdf_kind': pdf_kind,
            'total_chars': total_chars,
            'paragraphs': paras,
            'page_count': doc.page_count,
        }
    finally:
        doc.close()


# ---------------- OCR（图片型 PDF） ----------------

_engine_singleton = None
_engine_lock = threading.Lock()


def _ocr_engine():
    """懒加载单例：OCR 模型加载耗时数秒，复用避免每次重建"""
    global _engine_singleton
    if not OCR_AVAILABLE:
        raise PdfError('OCR 组件未安装（缺少 rapidocr_onnxruntime），无法识别扫描版 PDF')
    with _engine_lock:
        if _engine_singleton is None:
            _engine_singleton = RapidOCR()
        return _engine_singleton


def start_ocr(pdf_path, status_path):
    """后台线程：逐页渲染 + OCR 识别，进度与结果写入 status_path
    status JSON 结构：
      {running, page, total, error?, paragraphs?}
        running=True 时 paragraphs 为 null，page 表示已识别页数
        running=False 且有 paragraphs 表完成；有 error 表失败
    同一 status_path 上已有进行中/已完成任务时不重复启动。
    """
    if os.path.isfile(status_path):
        try:
            with open(status_path, encoding='utf-8') as f:
                st = json.load(f)
            if st.get('running') or st.get('paragraphs') is not None:
                return
        except Exception:
            pass

    def _write(st):
        try:
            with open(status_path, 'w', encoding='utf-8') as f:
                json.dump(st, f, ensure_ascii=False)
        except OSError:
            pass

    def worker():
        seq = 0
        paras = []
        try:
            doc = pymupdf.open(pdf_path)
            total = doc.page_count
            _write({'running': True, 'page': 0, 'total': total})
            engine = _ocr_engine()
            all_rows = []        # {text, h} 全部页聚类后的段落（后判标题）
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=OCR_DPI)
                png = '%s_p%d.png' % (status_path, i)
                pix.save(png)
                try:
                    rows = []
                    result, _ = engine(png)
                    if result:
                        for box, text, score in result:
                            t = str(text).strip()
                            if not t or len(t) > 800:
                                continue
                            if t.isdigit() or all(ch in " .,%+-*/=<>:;|!?()[]{}'\"\`~^#$@&" for ch in t):
                                continue
                            xs = [p[0] for p in box]
                            ys = [p[1] for p in box]
                            x0, x1 = min(xs), max(xs)
                            y0, y1 = min(ys), max(ys)
                            rows.append({'x0': x0, 'y0': y0, 'x1': x1, 'y1': y1,
                                         'h': y1 - y0 or 12, 'text': t})
                    # 页内聚类：行间距小 → 同段；空行 → 新段（物理行合并成段落）
                    clusters = []
                    for r in rows:
                        if clusters:
                            gap = r['y0'] - clusters[-1]['lines'][-1]['y1']
                            line_h = clusters[-1]['lines'][-1]['h']
                            if gap > max(line_h * 1.5, 8):
                                clusters.append({'lines': [r]})
                            else:
                                clusters[-1]['lines'].append(r)
                        else:
                            clusters.append({'lines': [r]})
                    for c in clusters:
                        # 段内行合并：CJK 相邻直连，其他补空格
                        text = ''
                        max_h = 0.0
                        boxes = []
                        for rrow in c['lines']:
                            lt = rrow['text']
                            if not lt:
                                continue
                            if text and not (_is_cjk(text[-1]) or _is_cjk(lt[0])):
                                text += ' '
                            text += lt
                            max_h = max(max_h, rrow['h'])
                            boxes.append([round(rrow['x0']), round(rrow['y0']),
                                          round(rrow['x1']), round(rrow['y1'])])
                        all_rows.append({'text': text, 'h': max_h, 'page': i, 'boxes': boxes})
                finally:
                    try:
                        os.remove(png)
                    except OSError:
                        pass
                _write({'running': True, 'page': i + 1, 'total': total})
            # 标题基准：行高中位数，显著更高且文本较短 → 标题
            heights = sorted(r['h'] for r in all_rows)
            med_h = heights[len(heights) // 2] if heights else 20
            for r in all_rows:
                is_title = r['h'] >= med_h * 1.3 and len(r['text']) <= TITLE_MAX_LEN
                paras.append({'seq': seq, 'original': r['text'], 'title': is_title,
                              'page': r['page'], 'boxes': r['boxes']})
                seq += 1
            doc.close()
            _write({'running': False, 'page': total, 'total': total, 'paragraphs': paras})
        except Exception as e:
            _write({'running': False, 'page': -1, 'total': 0, 'error': str(e), 'paragraphs': []})

    _write({'running': True, 'page': 0, 'total': 0})
    threading.Thread(target=worker, daemon=True).start()


def ocr_status(status_path):
    """读 OCR 进度；文件不存在视为未开始"""
    if os.path.isfile(status_path):
        try:
            with open(status_path, encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {'running': False, 'page': 0, 'total': 0, 'paragraphs': None}


def _page_main_bg(page):
    """整页主色（低清渲染 + 量化桶）：白底文档 → 白，深色页面 → 深色基准"""
    try:
        pix = page.get_pixmap(dpi=36)
        from collections import Counter as _C
        buckets = _C()
        w, h = pix.width, pix.height
        step = max(1, (w * h) // 400)
        n = 0
        for y in range(0, h, 2):
            for x in range(0, w, 2):
                if n % step == 0:
                    r_, g_, b_ = pix.pixel(x, y)[:3]
                    buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
                n += 1
        if not buckets:
            return (1.0, 1.0, 1.0)
        br, bg, bb = buckets.most_common(1)[0][0]
        return ((br * 16 + 8) / 255.0, (bg * 16 + 8) / 255.0, (bb * 16 + 8) / 255.0)
    except Exception:
        return (1.0, 1.0, 1.0)


def _rect_main_bg(page, rect):
    """rect 区域主色（低清渲染 + 量化桶）"""
    try:
        pix = page.get_pixmap(clip=pymupdf.Rect(*rect), dpi=36)
        from collections import Counter as _C
        buckets = _C()
        w, h = pix.width, pix.height
        if w < 1 or h < 1:
            return (1.0, 1.0, 1.0)
        step = max(1, (w * h) // 96)
        n = 0
        for y in range(0, h):
            for x in range(0, w):
                if n % step == 0:
                    r_, g_, b_ = pix.pixel(x, y)[:3]
                    buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
                n += 1
        if not buckets:
            return (1.0, 1.0, 1.0)
        br, bg, bb = buckets.most_common(1)[0][0]
        return ((br * 16 + 8) / 255.0, (bg * 16 + 8) / 255.0, (bb * 16 + 8) / 255.0)
    except Exception:
        return (1.0, 1.0, 1.0)


def _sample_bg(page, rect, page_bg):
    """估算文字框背景色：
    - 框内主色与页面主色接近（普通白底黑字）→ 直接用页面主色（稳，不受邻行污染）
    - 框内主色与页面主色差异大（深蓝卡片上的文字）→ 取框内主色（即深底色）
    返回 0~1 RGB。
    """
    try:
        inner = _rect_main_bg(page, rect)
        ir, ig, ib = inner
        pr, pg_, pb = page_bg
        diff = abs(ir - pr) + abs(ig - pg_) + abs(ib - pb)
        return page_bg if diff < 0.9 else inner
    except Exception:
        return page_bg


def _text_color(bg):
    """背景色 → 文字对比色：亮度高用黑字，亮度低用白字"""
    r, g, b = bg
    lum = 0.299 * r + 0.587 * g + 0.114 * b
    return (0.05, 0.05, 0.05) if lum >= 0.55 else (1.0, 1.0, 1.0)


def _layout_text(text, w, h, max_lines=3):
    """按框宽/高排布译文：尽量大字号、最多 max_lines 行，返回 (lines, fontsize)"""
    measure = _MEASURE_FONT.text_length

    def _wrap(t, fs, cap):
        out = []
        cur = ''
        for ch in t:
            if measure(cur + ch, fontsize=fs) > w - 2 and cur:
                out.append(cur)
                cur = ch
            else:
                cur += ch
            if len(out) >= cap:
                break
        if cur and len(out) < cap:
            out.append(cur)
        while out and len(out) > max_lines:
            out.pop()
        if out and measure(out[-1], fontsize=fs) > w - 2:
            s = out[-1]
            while s and measure(s + '…', fontsize=fs) > w - 2:
                s = s[:-1]
            out[-1] = s + '…'
        return out

    for fs in range(16, 4, -1):
        lines = _wrap(text, fs, 8)
        if not lines:
            continue
        if len(lines) <= max_lines and len(lines) * fs * 1.25 <= h:
            return lines, fs
    return _wrap(text, 5, max_lines), 5


def _pix_ring_bg(pix, boxes, pad=4):
    """从渲染图上采样段落框外侧环形带的背景主色（像素级，boxes 为像素坐标）"""
    try:
        from collections import Counter as _C
        buckets = _C()
        w, h = pix.width, pix.height
        for b in boxes:
            x0, y0, x1, y1 = b
            xl, xr = max(0, x0 - pad), min(w - 1, x1 + pad)
            yt, yb = max(0, y0 - pad), min(h - 1, y1 + pad)
            for x in range(x0 - pad, x0):
                if 0 <= x < w:
                    for y in (max(0, y0 - pad), min(h - 1, y1 + pad)):
                        r_, g_, b_ = pix.pixel(x, y)[:3]
                        buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
            for x in range(x1 + 1, x1 + 1 + pad):
                if 0 <= x < w:
                    for y in (max(0, y0 - pad), min(h - 1, y1 + pad)):
                        r_, g_, b_ = pix.pixel(x, y)[:3]
                        buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
            for y in range(y0 - pad, y0):
                if 0 <= y < h:
                    for x in (max(0, x0 - pad), min(w - 1, x1 + pad)):
                        r_, g_, b_ = pix.pixel(x, y)[:3]
                        buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
            for y in range(y1 + 1, y1 + 1 + pad):
                if 0 <= y < h:
                    for x in (max(0, x0 - pad), min(w - 1, x1 + pad)):
                        r_, g_, b_ = pix.pixel(x, y)[:3]
                        buckets[(r_ >> 4, g_ >> 4, b_ >> 4)] += 1
        if not buckets:
            return (1.0, 1.0, 1.0)
        br, bg, bb = buckets.most_common(1)[0][0]
        return ((br * 16 + 8) / 255.0, (bg * 16 + 8) / 255.0, (bb * 16 + 8) / 255.0)
    except Exception:
        return (1.0, 1.0, 1.0)


def export_pdf_overlay(src_bytes, paras, translated_map):
    """图片翻译式导出（OCR 型 PDF）：每页渲染成图 → 原图保持 → 在识别框上覆盖译文。
    特性：
      - 底图与原文版式 100% 一致（页面即图片）
      - 已翻译：框区填背景色 + 写译文（对比色）；未翻译：原样保留，不碰
      - 不依赖文字层/redact，扫描件与文字层乱码文档都适用
    paras 的 boxes 为 OCR 像素坐标（与 OCR_DPI 渲染一致）
    """
    doc = pymupdf.open(stream=src_bytes, filetype='pdf')
    out = pymupdf.open()
    try:
        scale = 72.0 / OCR_DPI       # 像素 -> pt
        for pno, src_page in enumerate(doc):
            pix = src_page.get_pixmap(dpi=OCR_DPI)
            page = out.new_page(width=src_page.rect.width, height=src_page.rect.height)
            page.insert_image(page.rect, pixmap=pix)
            for p in paras:
                if p.get('page', 0) != pno:
                    continue
                rects = p.get('boxes') or []
                if not rects:
                    continue
                text = (translated_map.get(p['seq']) or '').strip()
                if not text or text == p.get('original'):
                    continue        # 未翻译：原图原文保留
                x0 = min(r[0] for r in rects) * scale
                y0 = min(r[1] for r in rects) * scale
                x1 = max(r[2] for r in rects) * scale
                y1 = max(r[3] for r in rects) * scale
                w = x1 - x0
                h = y1 - y0
                if w < 4 or h < 4:
                    continue
                bg = _pix_ring_bg(pix, rects)
                page.draw_rect(pymupdf.Rect(x0 - 1, y0 - 1, x1 + 1, y1 + 1),
                               color=None, fill=bg)
                lines, fs = _layout_text(text, w, h)
                color = _text_color(bg)
                line_h = fs * 1.25
                ty = y0 + (h - line_h * len(lines)) / 2 + fs * 0.95
                for ln in lines:
                    tw = _MEASURE_FONT.text_length(ln, fontsize=fs)
                    page.insert_text((x0 + (w - tw) / 2, ty), ln,
                                     fontsize=fs, color=color,
                                     fontname=MAIN_FONT_NAME, fontfile=MAIN_FONT_FILE)
                    ty += line_h
        buf = io.BytesIO()
        out.subset_fonts()               # 字体子集化
        out.save(buf, garbage=3, deflate=True)
        return buf.getvalue(), 'translated.pdf', 'application/pdf'
    finally:
        doc.close()
        out.close()


def export_pdf_inline(src_bytes, paras, translated_map):
    """原排版导出：擦除原文（redact）→ 同坐标绘制译文。
    文本型：文字层被擦除、插图/背景保留；扫描件：整页图保留、识别框区域白底盖字后绘译文。
    paras: [{seq, original, title, page, rects|boxes}]；translated_map: {seq: 译文}
    """
    doc = pymupdf.open(stream=src_bytes, filetype='pdf')
    try:
        # ---- 1) 采样背景色 + 擦除原文（用背景色填充，深色卡片不露白块）----
        bg_cache = {}        # ((page, rect_tuple)) -> (r,g,b)
        page_bgs = {}        # page_index -> 整页主色
        # 只处理「有译文」的段落；未翻译（译文缺失/与原文相同）保留原文，不擦不写
        todo = []
        for p in paras:
            text = (translated_map.get(p['seq']) or '').strip()
            if not text or text == p.get('original'):
                continue
            rects = p.get('rects') or p.get('boxes') or []
            if not rects:
                continue
            todo.append((p, text, rects))
        for p, _text, rects in todo:
            pi = p.get('page', 0)
            pg = doc[pi]
            if pi not in page_bgs:
                page_bgs[pi] = _page_main_bg(pg)
            for r in rects:
                key = (pi, tuple(r))
                bg = _sample_bg(pg, r, page_bgs[pi])
                bg_cache[key] = bg
                # 外扩 2pt：字形可能溢出 bbox（CJK 字体常见），确保笔画完整擦除
                rr = pymupdf.Rect(*r)
                rr.x0 -= 2
                rr.y0 -= 2
                rr.x1 += 2
                rr.y1 += 2
                pg.add_redact_annot(rr, fill=bg)
        for pg in doc:
            pg.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE)   # 保护整页图（扫描件）
        # ---- 2) 写译文（颜色按背景反色）----
        for p, text, rects in todo:
            pg = doc[p.get('page', 0)]
            x0 = min(r[0] for r in rects)
            y0 = min(r[1] for r in rects)
            x1 = max(r[2] for r in rects)
            y1 = max(r[3] for r in rects)
            w = x1 - x0
            h = y1 - y0
            if w < 4 or h < 4:
                continue
            lines, fs = _layout_text(text, w, h)
            line_h = fs * 1.25
            ty = y0 + (h - line_h * len(lines)) / 2 + fs * 0.95
            # 背景色取该段首个框的采样，文字反色保证可读
            first_key = (p.get('page', 0), tuple(rects[0]))
            color = _text_color(bg_cache.get(first_key, (1.0, 1.0, 1.0)))
            for ln in lines:
                tw = _MEASURE_FONT.text_length(ln, fontsize=fs)
                tx = x0 + (w - tw) / 2
                pg.insert_text((tx, ty), ln, fontsize=fs, color=color,
                               fontname=MAIN_FONT_NAME, fontfile=MAIN_FONT_FILE)
                ty += line_h
        # ---- 3) 保存 ----
        buf = io.BytesIO()
        doc.subset_fonts()               # 字体子集化，避免微软雅黑整包进 PDF
        doc.save(buf, garbage=3, deflate=True)
        return buf.getvalue(), 'translated.pdf', 'application/pdf'
    finally:
        doc.close()


# ---------------- 导出 ----------------

def export_translated(original_paras, translated_paras, fmt, src_bytes=None):
    """按 fmt 生成译文文档，返回 (bytes, filename, mime)
    original_paras: [{seq|idx, original}]；translated_paras: [{idx, translated}]
    fmt: 'docx' | 'md' | 'txt' | 'pdf'
    """
    # 合并：按出现序，译文取 translated 的值（缺失则用原文）；保留标题层级标记
    by_idx = {}
    title_map = {}
    for i, p in enumerate(original_paras):
        by_idx[i] = p.get('original', '')
        title_map[i] = bool(p.get('title'))
    for tp in translated_paras:
        idx = tp.get('seq', tp.get('idx', 0))
        if tp.get('translated'):
            by_idx[idx] = tp['translated']
    ordered = [{'text': by_idx[i], 'title': title_map.get(i, False)} for i in sorted(by_idx)]

    if fmt == 'md':
        lines = []
        for item in ordered:
            lines.append(('## ' if item['title'] else '') + item['text'])
        text = '# 译文文档\n\n' + '\n\n'.join(lines) + '\n'
        return text.encode('utf-8'), 'translated.md', 'text/markdown'

    if fmt == 'txt':
        text = '\n'.join(item['text'] for item in ordered) + '\n'
        return text.encode('utf-8'), 'translated.txt', 'text/plain;charset=utf-8'

    if fmt == 'docx':
        wd = Document()
        wd.add_heading('译文文档', 0)
        for item in ordered:
            if item['title']:
                wd.add_heading(item['text'], level=1)
            else:
                para = wd.add_paragraph(item['text'])
                # 默认字体（正文小五→合适字号）
                for run in para.runs:
                    run.font.size = Pt(11)
        buf = io.BytesIO()
        wd.save(buf)
        return buf.getvalue(), 'translated.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    if fmt == 'pdf':
        # 原排版导出（需坐标）；无坐标时退回重排：
        #   - OCR 型（boxes 像素坐标）→ 整页渲染成图 + 图上覆盖（图片翻译式）
        #   - 文本型（rects 坐标）    → 文字层擦除 + 原位写入
        if src_bytes and original_paras and (original_paras[0].get('rects') or original_paras[0].get('boxes')):
            if original_paras[0].get('boxes') and not original_paras[0].get('rects'):
                return export_pdf_overlay(src_bytes, original_paras, by_idx)
            return export_pdf_inline(src_bytes, original_paras, by_idx)
        return _export_pdf_reflow(ordered)

    raise PdfError('不支持的导出格式: ' + fmt)


def _export_pdf_reflow(ordered):
    """重排导出：保持标题层级，版式从简（无坐标时回退用）"""
    out = pymupdf.open()
    fs = 11
    line_h = 17
    fnt = _MEASURE_FONT                 # 与渲染同一字体（微软雅黑或内置 CJK）
    x0 = 60
    max_x = 535                          # A4 595 - 右边距 60
    max_y = 792                          # A4 842 - 底边距 50
    pg = out.new_page()                  # 首页：标题 + 版式说明
    y = 50
    pg.insert_text((x0, y), '译文文档', fontsize=16, fontname=MAIN_FONT_NAME, fontfile=MAIN_FONT_FILE)
    y += 20
    pg.insert_text((x0, y), '此译文由 Universal Translator 重新排版生成，版式与原文可能存在差异',
                   fontsize=8, fontname=MAIN_FONT_NAME, fontfile=MAIN_FONT_FILE, color=(0.45, 0.45, 0.45))
    y += 26
    for item in ordered:
        text = item['text']
        title = item['title']
        # 标题：大字号 + 前后留空行；正文：11pt 连续排
        use_fs = 14 if title else fs
        if title:
            y += 8
        # 手动换行：按字体度量逐字符累积，超宽断行（中英混排都适用）
        chunks = []
        cur = ''
        for ch in text:
            if fnt.text_length(cur + ch, fontsize=use_fs) > (max_x - x0) and cur:
                chunks.append(cur)
                cur = ch
            else:
                cur += ch
        if cur:
            chunks.append(cur)
        for chunk in chunks:
            if y > max_y:
                pg = out.new_page()
                y = 60
            pg.insert_text((x0, y), chunk, fontsize=use_fs, fontname=MAIN_FONT_NAME, fontfile=MAIN_FONT_FILE)
            y += line_h + (2 if title else 0)
        if title:
            y += 6
    out.subset_fonts()               # 嵌入字体子集化：只保留用到的字形（微软雅黑 20MB -> 几 KB）
    buf = io.BytesIO()
    out.save(buf)
    out.close()
    return buf.getvalue(), 'translated.pdf', 'application/pdf'
